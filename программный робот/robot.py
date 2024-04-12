import webbrowser
import selenium
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import time
import openpyxl
import os
from bs4 import BeautifulSoup
import pandas as pd
from selenium.webdriver import Chrome
from openpyxl import Workbook
from docx import Document
import datetime
import win32com.client as client
driver = webdriver.Chrome()
url = 'https://licenses.roscosmos.ru/'
excel_file_path = 'C:/рита/Настройка_переодичности.xlsx'
workbook = openpyxl.open(excel_file_path, read_only=False)
sheet = workbook.active
time_interval = int(sheet.cell(row=2, column=2).value)
max_attempts = int(sheet.cell(row=2, column=2).value)
current_attempt = 1
mail_to = 'ilya.rusakov222@gmail.com'
outlook = client.Dispatch('Outlook.Application')
mail = outlook.CreateItem(0)
mail.To = mail_to

while current_attempt <= max_attempts:
    try:
        webbrowser.open(url)
        break

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        print(f"Повторная попытка {current_attempt}/{max_attempts}")
        current_attempt += 1
        time.sleep(time_interval)

else:
    print(f"Не удалось загрузить страницу за {max_attempts} попыток.")
if current_attempt == max_attempts:
    mail.Subject = 'Передача данных о лицензиях на космическую деятельность с сайта Роскосмос. Ошибка доступа к порталу'
    mail.Body = 'Доступ к порталу не был получен. Обратитесь в техническую поддержку интернет-портала. Данное письмо сформировано автоматически программным роботом, отвечать на него не нужно. При возникновении вопросов оформите, пожалуйста, обращение на портале самообслуживания'
    mail.Send()

# Закрываем книгу Excel после использования
workbook.close()
url = 'https://licenses.roscosmos.ru/'
driver = Chrome()
driver.get(url)
time.sleep(25)
mas = []
excel_file = openpyxl.open(
    'C:/рита/Шаблон_отчёта.xlsx', read_only=False)
df = pd.read_excel(
    'C:/рита/Шаблон_отчёта.xlsx')
sheet = excel_file.active

frame = driver.find_element(By.XPATH, '//app-table')
# driver.switch_to.frame(frame)

driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
time.sleep(10)

for page in range(1, 11):
    elem = driver.find_elements(
        By.XPATH, f"//ul[@class = 'pagination']//li{[page]}//a")
    x = driver.find_elements(
        By.XPATH, f"//a[@class = 'page-link nav-button']")[-1]
    x.click()
    # driver.find_element(
    #     By.XPATH, f"//a[@class = 'page-link nav-button']").click()
        
    time.sleep(10)
    table_data = driver.find_elements(
        # By.XPATH, f"//table[@class = 'search-result table table-hover']//tbody//tr//td")
        By.XPATH, f"//table[@class = 'table']//tbody//tr//td")
    table_data = [value for index, value in enumerate(
        table_data) if (index + 1) % 6 != 0 and (index+1) %3 !=0]
    # table_data = [value for index, value in enumerate(
    #     table_data) if (index + 1) % 3 != 0]
    for data in table_data:
        print(data.get_attribute("innerHTML"))
        mas.append(data.get_attribute("innerHTML"))

result = [mas[i:i + 4] for i in range(0, len(mas), 4)]
out_list = []
keys = ["Номер лицензии", "Наименование лицензиата",
        "ИНН", "Дата начала действия лицензии"]

current_datetime = datetime.date.today()
rows = 2
for stroka in result:
    sheet[f'B{rows}'] = stroka[0]
    sheet[f'C{rows}'] = stroka[1]
    sheet[f'D{rows}'] = stroka[2]
    sheet[f'E{rows}'] = stroka[3]
    sheet[f'F{rows}'] = current_datetime

    #excel_file.save('Шаблон_отчета.xlsx')
    excel_file.save(
        f'C:/рита/{current_datetime}.xlsx')
    rows += 1


workbook = openpyxl.load_workbook(
    f'C:/рита/{current_datetime}.xlsx')
worksheet = workbook.active
data_collection_date = worksheet['F2'].value
# Шаг 4: Считаем количество записей, начиная со строки 2
row_count = worksheet.max_row - 1  # вычитаем заголовок
# Шаг 5: Сохраняем значение и количество в файле Word
output_file_name = f"C:/рита//Отчет за {current_datetime}.docx"
word_document = Document()
par1 = word_document.add_paragraph(
    f'Сбор данных произведен {current_datetime}.')
par2 = word_document.add_paragraph(
    f'Найдено 100 лицензий на космическую деятельность.')
word_document.save(output_file_name)
outlook = client.Dispatch('Outlook.Application')
mail = outlook.CreateItem(0)
mail.Subject = f'Отчет за {current_datetime}'
mail.Body = 'Отчёт о лицензиях на космическую деятельность.\n\n\nДанное письмо сформировано автоматически программным роботом, отвечать на него не нужно. При возникновении вопросов оформите, пожалуйста, обращение на ilya.rusakov222@gmail.com или на портале самообслуживания.'
mail.To = 'ilya.rusakov222@gmail.com'

#Добавление вложения
excel = fr"C:/рита/{current_datetime}.xlsx"
mail.Attachments.Add(excel)
mail.Attachments.Add(output_file_name)
#Отправка письма
mail.Send()
